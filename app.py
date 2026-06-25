
import os
import uuid

from flask import Flask, render_template, request, jsonify
from werkzeug.utils import secure_filename
import PyPDF2

from engine import ResumeDocument, JobDescriptionDocument, MatchEngine

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 5 * 1024 * 1024  # 5 MB per file
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text


def extract_text_from_docx(file_path):
    try:
        import docx
    except ImportError:
        raise RuntimeError("python-docx is required to read .docx files")
    document = docx.Document(file_path)
    return "\n".join(p.text for p in document.paragraphs)


def read_file(file_path):
    ext = file_path.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_path)
    if ext == "docx":
        return extract_text_from_docx(file_path)
    if ext == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    raise ValueError("Unsupported file format")


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/analyze", methods=["POST"])
def analyze():
    if "resume" not in request.files or "job_description" not in request.files:
        return jsonify({"error": "Both resume and job description files are required."}), 400

    resume_file = request.files["resume"]
    job_file = request.files["job_description"]

    for f in (resume_file, job_file):
        if f.filename == "" or not allowed_file(f.filename):
            return jsonify({"error": f"Unsupported or missing file: {f.filename}"}), 400

    token = uuid.uuid4().hex
    resume_path = os.path.join(UPLOAD_DIR, f"{token}_resume_{secure_filename(resume_file.filename)}")
    job_path = os.path.join(UPLOAD_DIR, f"{token}_job_{secure_filename(job_file.filename)}")

    try:
        resume_file.save(resume_path)
        job_file.save(job_path)

        resume_text = read_file(resume_path)
        job_text = read_file(job_path)

        if not resume_text.strip() or not job_text.strip():
            return jsonify({"error": "Could not extract any text from one of the files."}), 422

        resume_doc = ResumeDocument(resume_text)
        job_doc = JobDescriptionDocument(job_text)
        result = MatchEngine(resume_doc, job_doc).analyze()

        return jsonify(result)

    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500

    finally:
        for p in (resume_path, job_path):
            if os.path.exists(p):
                os.remove(p)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
